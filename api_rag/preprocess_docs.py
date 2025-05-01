import os
import shutil
import re
import fitz
import pandas as pd
from docx import Document
from langchain_core.documents import Document as LangDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma

def extract_text_from_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    sections = []
    current_section = None

    def is_section_header(text):
        # Detecta títulos numerados como 1.0, 2.1, 3.2.1 etc
        return bool(re.match(r"^\d+(\.\d+)*\s", text.strip()))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if is_section_header(text):
            if current_section:
                sections.append(current_section)
            current_section = text
        else:
            if current_section:
                current_section += "\n" + text
            else:
                current_section = text

    if current_section:
        sections.append(current_section)

    # Tabelas ainda agrupadas como [TABELA] separadas
    table_blocks = []
    for table in doc.tables:
        table_text = "[TABELA]\n"
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            row_text = " | ".join(cells)
            table_text += f"- {row_text}\n"
        table_blocks.append(table_text.strip())

    return sections + table_blocks

def extract_text_from_excel(file_path):
    text = ""
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name)
        table_text = f"[PLANILHA]: {sheet_name}\n"
        for index, row in df.iterrows():
            row_text = []
            for col_name, value in row.items():
                row_text.append(f"{col_name.strip()}: {str(value).strip()}")
            table_text += "- " + " | ".join(row_text) + "\n"
        text += table_text + "\n"
    return [text]

def docling_virtual(text_blocks):
    # Agora cada bloco já é uma seção agrupada
    return text_blocks

def load_documents_from_folder(folder_path):
    docs = []
    total_encontrados = 0
    ignorados = 0
    erros = 0

    for root, _, files in os.walk(folder_path):
        for file in files:
            total_encontrados += 1

            if file.startswith("~$"):
                print(f"Ignorado arquivo temporário: {file}")
                ignorados += 1
                continue

            path = os.path.join(root, file)
            ext = file.lower().split('.')[-1]

            try:
                if ext == "pdf":
                    raw_text = extract_text_from_pdf(path)
                    structured_blocks = [raw_text]
                elif ext == "docx":
                    structured_blocks = extract_text_from_docx(path)
                elif ext in ["xlsx", "xls"]:
                    structured_blocks = extract_text_from_excel(path)
                else:
                    print(f"Ignorado arquivo com extensão não tratada: {file}")
                    ignorados += 1
                    continue

                rel_path = os.path.relpath(path, folder_path)

                for block in structured_blocks:
                    docs.append(LangDocument(
                        page_content=block.strip(),
                        metadata={"source": rel_path}
                    ))

            except Exception as e:
                print(f"Erro ao carregar {file}: {e}")
                erros += 1

    print("\n🧹 Resumo de carregamento:")
    print(f"🔎 Total de arquivos encontrados: {total_encontrados}")
    print(f"✅ Documentos carregados: {len(docs)}")
    print(f"🚫 Ignorados: {ignorados}")
    print(f"⚠️ Erros: {erros}\n")

    return docs

def main():
    print("Inicializando processamento...")
    # Ao criar um novo Chroma com persist_directory, ele limpa os dados antigos
    embedding = OllamaEmbeddings(model="nomic-embed-text", base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"))
    Chroma(embedding_function=embedding, persist_directory="./chroma_db")

    print("Carregando documentos...")
    documentos = load_documents_from_folder("./documentos")

    print(f"{len(documentos)} blocos carregados. Dividindo...")
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(documentos)

    print(f"{len(chunks)} chunks criados. Vetorizando...")
    embedding = OllamaEmbeddings(model="nomic-embed-text", base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"))
    vectorstore = Chroma.from_documents(
        chunks,
        embedding=embedding,
        persist_directory="./chroma_db"
    )

    print("✅ Vetorstore criado e salvo com sucesso!")

if __name__ == "__main__":
    main()
